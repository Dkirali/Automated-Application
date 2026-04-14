import os
import re
import shutil
import threading
import time
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

# Simple rate limiter: minimum seconds between consecutive LLM calls
_LLM_MIN_INTERVAL = 4.0  # seconds between LLM calls — prevents quota exhaustion on free tiers
_llm_last_call = 0.0
_llm_lock = threading.Lock()

RESUMES_DIR = Path(__file__).parent.parent / "resumes"

CANDIDATE_HEADER = {
    "name":     "Doruk Kirali",
    "phone":    "0532 286 04 61",
    "email":    "kiralidoruk@gmail.com",
    "linkedin": "linkedin.com/in/doruk-kirali",
    "github":   "github.com/Dkirali",
}

# Lines containing these strings are stripped from LLM body output to avoid
# duplicating the fixed header we prepend.
_CONTACT_SKIP_PATTERNS = [
    "kiralidoruk@gmail.com", "0532", "doruk-kirali",
    "dkirali", "linkedin.com/in/doruk", "github.com/dkirali",
]

FIT_PROMPT = """You are a senior recruiter evaluating a candidate's fit for a role.

Job Posting:
{job_description}

Candidate Resume:
{resume_text}

Respond in this exact format (no extra text):
FIT_SCORE: <0-100>
STRENGTHS: <comma-separated list of 2-4 matching skills or experiences>
GAPS: <comma-separated list of 1-3 missing skills, or "None">
VERDICT: <one sentence — would you recommend applying? why?>
JD_SUMMARY: <2-3 sentence summary of the role, seniority level, and key focus areas>
JD_KEYWORDS: keyword1, keyword2, keyword3, ... <8-12 most important ATS/skills keywords from the job posting>"""

TAILOR_PROMPT = """You are an expert resume writer specialising in ATS optimisation.

Job Posting:
{job_description}

Current Resume:
{resume_text}

Task:
1. Extract the 8-12 most important ATS keywords from the job posting (skills, tools, methodologies, titles).
2. Rewrite the resume experience bullet points to incorporate these keywords where truthful.
   CRITICAL RULES for keyword incorporation:
   - Keywords must appear as PART OF THE NATURAL SENTENCE describing what was done. Good: "Led sprint planning using Agile methodologies". Bad: "Led sprint planning, demonstrating Agile skills".
   - NEVER append keyword labels to the end of a sentence (e.g. "showcasing Operational Efficiency", "utilizing Stakeholder Management", "demonstrating Communication skills"). This reads as spam.
   - NEVER start or end a bullet with a generic phrase like "applying X and Y skills" or "ensuring effective Z". Instead, describe the ACTUAL WORK that used those skills.
   - Each bullet should describe a concrete action and result, not list skill categories.
   - If a keyword cannot be naturally woven into existing content, place it in the SKILLS section instead.
3. Do NOT invent experience. Only rephrase existing content to better match the posting.
4. Format the RESUME block as plain text (no Markdown, no HTML):
   - Section headings in ALL CAPS on their own line (e.g. PROFESSIONAL EXPERIENCE, EDUCATION, SKILLS, CERTIFICATES)
   - Each role: "Company Name | Location | Start–End" on one line, then the job title on the next line
   - Use • for main bullet points, - for sub-bullets
   - Do NOT include the candidate name, phone, email, LinkedIn, or GitHub — those are added separately
5. In the CERTIFICATES section, always include these two courses if not already present:
   - Udemy — AI Coder: Vibe Coder to Agentic Engineer in 3 Weeks
   - Udemy — AI Engineer Agentic Track: The Complete Agent & MCP Course
6. You MAY add new bullet points to existing experience roles if the job posting requires a skill the candidate demonstrably has from their history. Do NOT invent entirely new roles or technologies not present anywhere in the resume.
7. Only use the keyword "AI" where it genuinely refers to artificial intelligence concepts. Do NOT treat words like "main", "rain", or "training" as containing the keyword "AI".

Respond in this exact format:
KEYWORDS: keyword1, keyword2, keyword3, ...
RESUME:
[Full rewritten resume as plain text, starting directly with the first section heading]"""

TAILOR_RETRY_PROMPT = """The following tailored resume scored below 80% ATS keyword match.

These keywords are MISSING from the resume and MUST be incorporated: {missing_keywords}

Job Posting:
{job_description}

Current Resume:
{resume}

Rewrite the resume to naturally incorporate ALL of the missing keywords listed above.
Keep all existing content and structure. Only add or rephrase content — do not remove anything.
Do NOT include contact info (name, phone, email, LinkedIn, GitHub).

Respond in this exact format:
RESUME:
[Full rewritten resume as plain text]"""


def _matches_keyword(kw: str, text: str) -> bool:
    """Check if a keyword appears in text using whole-word matching for short keywords."""
    kw_lower = kw.lower()
    text_lower = text.lower()
    if len(kw_lower) <= 4:
        return bool(re.search(r'\b' + re.escape(kw_lower) + r'\b', text_lower))
    return kw_lower in text_lower


def extract_keywords_from_response(response_text: str) -> list[str]:
    match = re.search(r"KEYWORDS:\s*([^\n]+)", response_text)
    if not match:
        return []
    return [k.strip() for k in match.group(1).split(",") if k.strip()]


def extract_resume_from_response(response_text: str) -> str:
    match = re.search(r"RESUME:\s*\n(.*)", response_text, re.DOTALL)
    return match.group(1).strip() if match else ""


def calculate_ats_score(keywords: list[str], resume_text: str) -> int:
    if not keywords:
        return 0
    matched = sum(1 for kw in keywords if _matches_keyword(kw, resume_text))
    return round((matched / len(keywords)) * 100)


def calculate_original_ats_score(keywords: list[str], master_resume_path: str) -> int:
    """Compute ATS score against the unmodified master resume."""
    master_path = Path(master_resume_path)
    if not master_path.exists() or not keywords:
        return 0
    original_text = read_resume_text(master_path)
    return calculate_ats_score(keywords, original_text)


def strip_markdown(text: str) -> str:
    """Remove Markdown markers before writing to .docx (keeps download clean)."""
    text = re.sub(r'^#{1,3}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    return text


ALLOWED_EXTENSIONS = {".docx", ".doc", ".pdf"}


def read_resume_text(path: Path) -> str:
    """Read plain text from .docx, .doc, or .pdf resume files."""
    ext = path.suffix.lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif ext == ".doc":
        import mammoth
        with open(path, "rb") as f:
            result = mammoth.extract_raw_text(f)
        return result.value.strip()
    elif ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(
                page.extract_text() or "" for page in pdf.pages
            ).strip()
    else:
        raise ValueError(f"Unsupported resume format: {ext}")


# Keep old name as alias so detail view route still works
def read_docx_text(path: Path) -> str:
    return read_resume_text(path)


def _add_section_heading(doc, text: str):
    """Add a bold ALL-CAPS section heading with a bottom border rule."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


SECTION_KEYWORDS = {
    "professional experience", "experience", "education", "skills", "hard skills",
    "management skills", "about", "about me", "summary", "references", "contact",
    "projects", "certifications", "certificates", "languages", "tools", "internships",
}


def write_tailored_docx(master_path: Path, new_text: str, output_path: Path):
    """Write tailored resume to .docx with fixed header block."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if master_path.suffix.lower() == ".docx":
        # For .docx masters, still inject header at the start
        shutil.copy(master_path, output_path)
        doc = Document(output_path)
        # Clear all existing paragraphs and rebuild
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = ""
        doc.save(output_path)
        # Fall through to rebuild cleanly
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Pt(36)
            section.left_margin = section.right_margin = Pt(54)
    else:
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Pt(36)
            section.left_margin = section.right_margin = Pt(54)

    # ── Fixed header ──────────────────────────────────────────────────────────
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CANDIDATE_HEADER["name"])
    r.bold = True
    r.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(2)

    # Job title (last held role — fixed)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Product Operations Manager")
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    # Contact line
    contact = (
        f"{CANDIDATE_HEADER['phone']} • {CANDIDATE_HEADER['email']} • "
        f"{CANDIDATE_HEADER['linkedin']} • {CANDIDATE_HEADER['github']}"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(contact)
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(8)

    # ── Body ─────────────────────────────────────────────────────────────────
    # Detect lines with pipe separators as company/role metadata:
    #   "Company Name | Location | Dates"  → company line
    #   Next non-bullet line after a company line → job title (bold)
    lines = new_text.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        i += 1
        if not stripped:
            continue

        # Skip contact duplicates
        low = stripped.lower()
        if any(pat in low for pat in _CONTACT_SKIP_PATTERNS):
            continue
        if stripped.lower() == CANDIDATE_HEADER["name"].lower():
            continue

        # Section heading
        if low.rstrip(":") in SECTION_KEYWORDS:
            _add_section_heading(doc, stripped)
            continue

        # Bullet points
        if stripped.startswith(("•", "-", "*")):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped.lstrip("•-* "))
            p.paragraph_format.space_after = Pt(1)
            continue

        # Company/role line: contains | separators (e.g. "Styx Intelligence | Vancouver, Canada | 2022–2024")
        if "|" in stripped:
            parts = [p.strip() for p in stripped.split("|")]
            company = parts[0]
            location_dates = " | ".join(parts[1:])

            # Add space before each new experience entry
            p_spacer = doc.add_paragraph()
            p_spacer.paragraph_format.space_before = Pt(4)
            p_spacer.paragraph_format.space_after = Pt(0)

            # Company name — regular weight on its own line
            p = doc.add_paragraph()
            r = p.add_run(company)
            r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

            # Next non-empty line should be the job title
            title_line = ""
            while i < len(lines):
                candidate = lines[i].strip()
                i += 1
                if candidate:
                    title_line = candidate
                    break

            # Job title (bold) with location | dates right-aligned
            if title_line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                # Use a tab stop at the right margin for right-alignment
                from docx.shared import Inches
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9360')  # ~6.5 inches
                tabs.append(tab)
                pPr.append(tabs)

                r = p.add_run(title_line)
                r.bold = True
                r.font.size = Pt(10)
                if location_dates:
                    r2 = p.add_run(f"\t{location_dates}")
                    r2.font.size = Pt(10)
            continue

        # Any other plain text line
        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_after = Pt(2)

    doc.save(output_path)


def export_pdf(docx_path: Path, pdf_path: Path):
    import mammoth
    from weasyprint import HTML
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)
    HTML(string=result.value).write_pdf(str(pdf_path))


def _call_provider(name: str, fn, logger) -> str | None:
    """Try a single LLM provider with up to 3 retries and exponential backoff.
    Rate-limit (429/quota) errors skip immediately to the next provider.
    Returns the response text on success, or None on failure."""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            return fn()
        except Exception as e:
            err_str = str(e).lower()
            is_rate_limit = "429" in err_str or "rate" in err_str or "quota" in err_str or "resource_exhausted" in err_str
            if is_rate_limit:
                # Don't waste time retrying — skip to next provider immediately
                logger.info("%s rate-limited, skipping to next provider: %s", name, e)
                return None
            wait = 2 * (2 ** attempt)  # 2s, 4s, 8s
            is_last = attempt == max_retries - 1
            if is_last:
                logger.warning("%s failed after %d attempts: %s", name, max_retries, e)
                return None
            logger.info("%s attempt %d failed, retrying in %ds: %s",
                        name, attempt + 1, wait, e)
            time.sleep(wait)
    return None


# All available models the user can pick from
AVAILABLE_MODELS = {
    "auto":                         "Auto (best available)",
    "groq/llama-3.3-70b":          "Groq — Llama 3.3 70B",
    "openrouter/gpt-oss-120b":     "OpenRouter — GPT-OSS 120B",
    "openrouter/minimax-m2.5":     "OpenRouter — MiniMax M2.5",
    "openrouter/free":             "OpenRouter — Best Free",
    "anthropic/claude-sonnet":     "Anthropic — Claude Sonnet",
}

# Last model that successfully produced a result (module-level for display)
_last_model_used = "—"
_last_model_lock = threading.Lock()


def get_last_model_used() -> str:
    with _last_model_lock:
        return _last_model_used


def _set_last_model(name: str):
    global _last_model_used
    with _last_model_lock:
        _last_model_used = name


def _track_usage(model_key: str) -> None:
    """Increment API usage counter. Lazy import to avoid circular dependency."""
    try:
        from db.database import increment_api_usage
        increment_api_usage(model_key)
    except Exception:
        pass  # non-critical — never let tracking failures affect tailoring


def _call_llm(prompt: str, max_tokens: int = 2048, preferred_model: str = "auto") -> str:
    """Call an LLM provider.
    If preferred_model is set (not 'auto'), call only that provider.
    Otherwise cascade: Groq → OpenRouter (3 models) → Anthropic.
    Rate-limited to at most one call every _LLM_MIN_INTERVAL seconds."""
    import logging
    _logger = logging.getLogger("jobbot.resume")

    # Enforce minimum interval between LLM calls to avoid quota exhaustion
    global _llm_last_call
    wait = 0.0
    with _llm_lock:
        elapsed = time.time() - _llm_last_call
        if elapsed < _LLM_MIN_INTERVAL:
            wait = _LLM_MIN_INTERVAL - elapsed
        _llm_last_call = time.time() + wait  # reserve our slot
    if wait > 0:
        _logger.debug("Rate limiter: waiting %.1fs before next LLM call", wait)
        time.sleep(wait)

    groq_key = os.environ.get("GROQ_API_KEY")
    openrouter_key = os.environ.get("OPENROUTER_API_KEY")
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY")

    errors = []

    # ── Provider helpers ─────────────────────────────────────────────────
    def _groq():
        from groq import Groq
        client = Groq(api_key=groq_key)
        message = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.choices[0].message.content

    def _or_call(model_id):
        import requests as _req
        resp = _req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {openrouter_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost:5001",
                "X-Title": "JobBot",
            },
            json={
                "model": model_id,
                "max_tokens": max_tokens,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=90,
        )
        resp.raise_for_status()
        data = resp.json()
        content = data["choices"][0]["message"]["content"]
        if not content:
            raise RuntimeError(f"Empty response from {model_id}")
        return content

    def _anthropic():
        import anthropic
        client = anthropic.Anthropic(api_key=anthropic_key)
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text

    # ── Specific model requested ─────────────────────────────────────────
    if preferred_model and preferred_model != "auto":
        _MODEL_MAP = {
            "groq/llama-3.3-70b":      ("Groq/Llama-3.3-70B", _groq, groq_key),
            "openrouter/gpt-oss-120b":  ("OpenRouter/gpt-oss-120b:free",  lambda: _or_call("openai/gpt-oss-120b:free"), openrouter_key),
            "openrouter/minimax-m2.5":  ("OpenRouter/minimax-m2.5:free",  lambda: _or_call("minimax/minimax-m2.5:free"), openrouter_key),
            "openrouter/free":          ("OpenRouter/free",               lambda: _or_call("openrouter/free"), openrouter_key),
            "anthropic/claude-sonnet":  ("Anthropic/Claude-Sonnet",       _anthropic, anthropic_key),
        }
        entry = _MODEL_MAP.get(preferred_model)
        if entry:
            label, fn, key_present = entry
            if key_present:
                result = _call_provider(label, fn, _logger)
                if result is not None:
                    _set_last_model(label)
                    _track_usage(preferred_model)
                    return result
            raise RuntimeError(f"Model {preferred_model} failed — no API key or provider error")
        # Unknown model, fall through to auto
        _logger.warning("Unknown preferred_model=%s, falling back to auto", preferred_model)

    # ── Auto cascade ─────────────────────────────────────────────────────
    if groq_key:
        result = _call_provider("Groq/Llama-3.3-70B", _groq, _logger)
        if result is not None:
            _set_last_model("Groq/Llama-3.3-70B")
            _track_usage("groq/llama-3.3-70b")
            return result
        errors.append("Groq")

    if openrouter_key:
        _OR_MODELS = [
            ("openai/gpt-oss-120b:free",  "OpenRouter/gpt-oss-120b:free",  "openrouter/gpt-oss-120b"),
            ("minimax/minimax-m2.5:free",  "OpenRouter/minimax-m2.5:free", "openrouter/minimax-m2.5"),
            ("openrouter/free",            "OpenRouter/free",               "openrouter/free"),
        ]
        for or_id, or_label, or_key in _OR_MODELS:
            result = _call_provider(or_label, lambda m=or_id: _or_call(m), _logger)
            if result is not None:
                _set_last_model(or_label)
                _track_usage(or_key)
                return result
            errors.append(or_label)

    if anthropic_key:
        result = _call_provider("Anthropic/Claude-Sonnet", _anthropic, _logger)
        if result is not None:
            _set_last_model("Anthropic/Claude-Sonnet")
            _track_usage("anthropic/claude-sonnet")
            return result
        errors.append("Anthropic")

    if errors:
        raise RuntimeError(f"All LLM providers failed: {', '.join(errors)}")
    raise RuntimeError("No API key set — add a Groq, OpenRouter, or Anthropic key in Settings")


def tailor_resume(job_id: int, job_description: str, master_resume_path: str,
                  existing_keywords: list[str] | None = None,
                  preferred_model: str = "auto") -> dict:
    master_path = Path(master_resume_path)
    if not master_path.exists():
        raise FileNotFoundError(f"Master resume not found: {master_resume_path}")
    resume_text = read_resume_text(master_path)

    try:
        response_text = _call_llm(TAILOR_PROMPT.format(
            job_description=job_description,
            resume_text=resume_text,
        ), max_tokens=2048, preferred_model=preferred_model)
    except Exception as e:
        raise RuntimeError(f"LLM API error: {e}") from e

    # Reuse stored keywords on re-tailor so the original score stays stable
    if existing_keywords:
        keywords = existing_keywords
    else:
        keywords = extract_keywords_from_response(response_text)

    tailored_text = extract_resume_from_response(response_text)
    if not tailored_text:
        tailored_text = resume_text  # fallback: use original if parse failed

    # Score against tailored text and original master
    ats_score = calculate_ats_score(keywords, tailored_text)
    original_ats_score = calculate_original_ats_score(keywords, master_resume_path)

    job_dir = RESUMES_DIR / str(job_id)
    job_dir.mkdir(parents=True, exist_ok=True)

    docx_path = job_dir / "tailored.docx"
    pdf_path = job_dir / "tailored.pdf"

    # Strip any residual Markdown before writing
    write_tailored_docx(master_path, strip_markdown(tailored_text), docx_path)

    try:
        export_pdf(docx_path, pdf_path)
    except Exception:
        pdf_path = None

    return {
        "keywords": keywords,
        "keywords_str": ", ".join(keywords),
        "ats_score": ats_score,
        "original_ats_score": original_ats_score,
        "tailored_text": tailored_text,
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path) if pdf_path else None,
        "model_used": get_last_model_used(),
    }


def parse_fit_score(text: str) -> int:
    m = re.search(r"FIT_SCORE:\s*(\d+)", text)
    return int(m.group(1)) if m else 0


def parse_fit_field(text: str, field: str) -> str:
    m = re.search(rf"{field}:\s*([^\n]+)", text)
    return m.group(1).strip() if m else ""


def generate_fit_summary(job_description: str, master_resume_path: str) -> dict:
    """
    Call LLM to evaluate how well the candidate fits the job.
    Returns: {fit_score, strengths, gaps, verdict, raw}
    """
    master_path = Path(master_resume_path)
    if not master_path.exists():
        raise FileNotFoundError(f"Master resume not found: {master_resume_path}")
    resume_text = read_resume_text(master_path)

    try:
        raw = _call_llm(FIT_PROMPT.format(
            job_description=job_description[:4000],
            resume_text=resume_text[:3000],
        ), max_tokens=512)
    except Exception as e:
        raise RuntimeError(f"LLM API error: {e}") from e

    return {
        "fit_score":    parse_fit_score(raw),
        "strengths":    [s.strip() for s in parse_fit_field(raw, "STRENGTHS").split(",") if s.strip()],
        "gaps":         [g.strip() for g in parse_fit_field(raw, "GAPS").split(",") if g.strip()],
        "verdict":      parse_fit_field(raw, "VERDICT"),
        "jd_summary":   parse_fit_field(raw, "JD_SUMMARY"),
        "jd_keywords":  parse_fit_field(raw, "JD_KEYWORDS"),
        "raw":          raw,
    }
